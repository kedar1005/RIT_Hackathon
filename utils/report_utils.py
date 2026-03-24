import pandas as pd
import io

def generate_word_report(df: pd.DataFrame) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.section import WD_ORIENT
    
    doc = Document()
    
    # Set to Landscape for better table fitting
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    doc.add_heading('Citizen AI - Complaints Report', 0)
    
    if df.empty:
        doc.add_paragraph("No data available for the selected filters.")
    else:
        # Convert all to string
        df = df.astype(str)
        
        # Create a single table, one row per complaint
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = str(column).replace('_', ' ').title()
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            
        # Data Rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
                
    # --- ADD CHARTS ---
    doc.add_page_break()
    doc.add_heading('Analytics Dashboard', 1)
    
    try:
        from utils.data_utils import get_complaint_stats, get_daily_trend, get_agent_leaderboard
        from ml.model_tracker import get_category_distribution_chart, get_urgency_donut, get_daily_trend_chart, get_agent_leaderboard_chart
        
        stats = get_complaint_stats()
        
        # 1. Category Chart
        by_cat = stats.get('by_category', [])
        if by_cat:
            fig_cat = get_category_distribution_chart(by_cat)
            img_cat = fig_cat.to_image(format="png", width=800, height=500)
            doc.add_heading('Complaints by Category', 2)
            doc.add_picture(io.BytesIO(img_cat), width=Inches(6.0))
            
        # 2. Urgency Donut
        high = stats.get('urgency_high', 0)
        medium = stats.get('urgency_medium', 0)
        low = stats.get('urgency_low', 0)
        if high > 0 or medium > 0 or low > 0:
            fig_urg = get_urgency_donut(high, medium, low)
            img_urg = fig_urg.to_image(format="png", width=800, height=500)
            doc.add_heading('Urgency Distribution', 2)
            doc.add_picture(io.BytesIO(img_urg), width=Inches(6.0))
        
        # 3. Daily Trend
        daily = get_daily_trend(30)
        if daily:
            fig_trend = get_daily_trend_chart(daily)
            img_trend = fig_trend.to_image(format="png", width=800, height=500)
            doc.add_heading('Daily Trend (30 days)', 2)
            doc.add_picture(io.BytesIO(img_trend), width=Inches(6.0))
            
        # 4. Agent Leaderboard
        agents = get_agent_leaderboard()
        if agents:
            fig_agents = get_agent_leaderboard_chart(agents)
            img_agents = fig_agents.to_image(format="png", width=800, height=500)
            doc.add_heading('Agent Leaderboard', 2)
            doc.add_picture(io.BytesIO(img_agents), width=Inches(6.0))
            
    except Exception as e:
        doc.add_paragraph(f"Could not load charts: {str(e)}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def generate_pdf_report(df: pd.DataFrame) -> bytes:
    from fpdf import FPDF
    
    pdf = FPDF(orientation="Landscape")
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 10, "Citizen AI - Complaints Report", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)
    
    if df.empty:
        pdf.set_font("helvetica", "", 12)
        pdf.cell(0, 10, "No data available for the selected filters.")
    else:
        pdf.set_font("helvetica", "B", 9)
        priority_cols = ['id', 'category', 'ai_urgency', 'status', 'created_at', 'department']
        cols = [c for c in priority_cols if c in df.columns]
        if not cols:
            cols = list(df.columns)[:6]
            
        col_width = 280 / len(cols)
        
        for col in cols:
            pdf.cell(col_width, 10, str(col).upper(), border=1)
        pdf.ln()
        
        pdf.set_font("helvetica", "", 8)
        for _, row in df.iterrows():
            for col in cols:
                val = str(row[col]).replace('\n', ' ')
                val = val.encode('latin-1', 'replace').decode('latin-1')
                if len(val) > 35:
                    val = val[:32] + "..."
                pdf.cell(col_width, 8, val, border=1)
            pdf.ln()
            
    # --- ADD CHARTS ---
    pdf.add_page()
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 10, "Analytics Dashboard", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)
    
    try:
        from utils.data_utils import get_complaint_stats, get_daily_trend, get_agent_leaderboard
        from ml.model_tracker import get_category_distribution_chart, get_urgency_donut, get_daily_trend_chart, get_agent_leaderboard_chart
        
        stats = get_complaint_stats()
        
        # 1. Category Chart
        by_cat = stats.get('by_category', [])
        if by_cat:
            fig_cat = get_category_distribution_chart(by_cat)
            img_cat = fig_cat.to_image(format="png", width=1200, height=700)
            pdf.set_font("helvetica", "B", 12)
            pdf.cell(0, 10, "Complaints by Category", new_x="LMARGIN", new_y="NEXT")
            pdf.image(io.BytesIO(img_cat), w=250)
            pdf.ln(10)
            
        # 2. Urgency Donut
        high = stats.get('urgency_high', 0)
        medium = stats.get('urgency_medium', 0)
        low = stats.get('urgency_low', 0)
        if high > 0 or medium > 0 or low > 0:
            fig_urg = get_urgency_donut(high, medium, low)
            img_urg = fig_urg.to_image(format="png", width=1200, height=700)
            pdf.add_page()
            pdf.set_font("helvetica", "B", 12)
            pdf.cell(0, 10, "Urgency Distribution", new_x="LMARGIN", new_y="NEXT")
            pdf.image(io.BytesIO(img_urg), w=250)
            pdf.ln(10)
        
        # 3. Daily Trend
        daily = get_daily_trend(30)
        if daily:
            fig_trend = get_daily_trend_chart(daily)
            img_trend = fig_trend.to_image(format="png", width=1200, height=700)
            pdf.add_page()
            pdf.set_font("helvetica", "B", 12)
            pdf.cell(0, 10, "Daily Trend (30 days)", new_x="LMARGIN", new_y="NEXT")
            pdf.image(io.BytesIO(img_trend), w=250)
            pdf.ln(10)
            
        # 4. Agent Leaderboard
        agents = get_agent_leaderboard()
        if agents:
            fig_agents = get_agent_leaderboard_chart(agents)
            img_agents = fig_agents.to_image(format="png", width=1200, height=700)
            pdf.add_page()
            pdf.set_font("helvetica", "B", 12)
            pdf.cell(0, 10, "Agent Leaderboard", new_x="LMARGIN", new_y="NEXT")
            pdf.image(io.BytesIO(img_agents), w=250)
            
    except Exception as e:
        pdf.set_font("helvetica", "", 10)
        pdf.cell(0, 10, f"Could not load charts: {str(e)}")

    return bytes(pdf.output())
